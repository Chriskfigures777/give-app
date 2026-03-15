"""
Generate The Exchange App — Facebook Video Ads Creative Brief Document
Creates a fully formatted .docx file with embedded Pexels images, ad copy, pain points, and CTAs.
Compatible with Microsoft Word and Apple Pages.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMAGES_DIR = os.path.join(BASE_DIR, "ad-assets", "images")
OUTPUT_PATH = os.path.join(BASE_DIR, "The Exchange App - Facebook Video Ads Creative Brief.docx")

# ── Color palette ──────────────────────────────────────────────────────────────
DARK_BG       = RGBColor(0x12, 0x14, 0x1A)   # near-black
BRAND_BLUE    = RGBColor(0x1A, 0x73, 0xE8)   # Exchange blue
BRAND_GOLD    = RGBColor(0xD4, 0xA0, 0x17)   # warm gold accent
PAIN_RED      = RGBColor(0xC0, 0x39, 0x2B)   # pain-point red
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY    = RGBColor(0xF5, 0xF5, 0xF5)
MID_GRAY      = RGBColor(0x88, 0x88, 0x88)
DARK_GRAY     = RGBColor(0x33, 0x33, 0x33)
GREEN_CTA     = RGBColor(0x1E, 0x8B, 0x4C)   # CTA green

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_paragraph_with_border(doc, text, font_size=11, bold=False, italic=False,
                               color=None, bg_hex=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                               space_before=0, space_after=6, left_indent=0):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left_indent:
        pf.left_indent = Pt(left_indent)

    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_divider(doc, color_hex="1A73E8"):
    """Horizontal rule divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_label_pill(doc, label: str, bg_hex: str, text_color=WHITE):
    """Colored label pill (e.g. 'VIDEO AD', 'PAIN POINT')."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    cell.width = Inches(2.2)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.font.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = text_color
    # Remove cell borders
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tcBorders.append(el)
    tcPr.append(tcBorders)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_ad_section(doc, ad_num, title, pain_hook, body_copy, cta_text, cta_subtext,
                   platform_notes, video_script_notes, image_path, image_caption,
                   ad_type_label="VIDEO AD", pain_label="PAIN POINT HOOK",
                   phase="PHASE 1 — COLD TRAFFIC", bg_hex_title="1A2340"):
    """Render a complete ad section with image, copy, and metadata."""

    # ── Section header bar ──────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_hex_title)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(8)
    r1 = p.add_run(f"  AD {ad_num:02d}  |  ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = BRAND_GOLD
    r2 = p.add_run(title.upper())
    r2.font.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = WHITE
    r3 = p.add_run(f"   ·   {phase}")
    r3.font.size = Pt(8)
    r3.font.color.rgb = MID_GRAY
    # remove table borders
    for row in tbl.rows:
        for c in row.cells:
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'nil')
                tcBorders.append(el)
            tcPr.append(tcBorders)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Two-column layout: image (left) + copy (right) ──────────────────────
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl2.columns[0].width = Inches(3.1)
    tbl2.columns[1].width = Inches(4.0)

    # Left cell: image
    img_cell = tbl2.rows[0].cells[0]
    img_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_bg(img_cell, "F0F2F5")

    ip = img_cell.paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_before = Pt(6)
    ip.paragraph_format.space_after = Pt(4)

    if image_path and os.path.exists(image_path):
        run_img = ip.add_run()
        run_img.add_picture(image_path, width=Inches(2.85))
    else:
        ip.add_run("[Image not found]").font.color.rgb = PAIN_RED

    # Caption under image
    cap_p = img_cell.add_paragraph(image_caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(2)
    cap_p.paragraph_format.space_after = Pt(6)
    cap_r = cap_p.runs[0] if cap_p.runs else cap_p.add_run(image_caption)
    cap_r.font.size = Pt(7.5)
    cap_r.font.italic = True
    cap_r.font.color.rgb = MID_GRAY

    # Ad type badge
    badge_p = img_cell.add_paragraph()
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_p.paragraph_format.space_before = Pt(0)
    badge_p.paragraph_format.space_after = Pt(8)
    br = badge_p.add_run(f"  {ad_type_label}  ")
    br.font.bold = True
    br.font.size = Pt(8)
    br.font.color.rgb = WHITE
    # Simulate badge with highlight
    rPr = br._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'darkBlue')
    rPr.append(highlight)

    # Right cell: ad copy
    copy_cell = tbl2.rows[0].cells[1]
    copy_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_bg(copy_cell, "FFFFFF")

    # Pain label
    pain_lp = copy_cell.add_paragraph()
    pain_lp.paragraph_format.space_before = Pt(8)
    pain_lp.paragraph_format.space_after = Pt(2)
    pain_lp.paragraph_format.left_indent = Pt(8)
    pr = pain_lp.add_run(f"▶  {pain_label}")
    pr.font.bold = True
    pr.font.size = Pt(7.5)
    pr.font.color.rgb = PAIN_RED

    # Pain hook text
    hook_p = copy_cell.add_paragraph()
    hook_p.paragraph_format.space_before = Pt(2)
    hook_p.paragraph_format.space_after = Pt(10)
    hook_p.paragraph_format.left_indent = Pt(8)
    hr = hook_p.add_run(pain_hook)
    hr.font.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = DARK_GRAY

    # Body copy label
    body_label_p = copy_cell.add_paragraph()
    body_label_p.paragraph_format.space_before = Pt(0)
    body_label_p.paragraph_format.space_after = Pt(3)
    body_label_p.paragraph_format.left_indent = Pt(8)
    blr = body_label_p.add_run("AD BODY COPY  (paste into Facebook Ads Manager)")
    blr.font.bold = True
    blr.font.size = Pt(7.5)
    blr.font.color.rgb = BRAND_BLUE

    # Body copy box
    body_tbl = copy_cell.add_table(rows=1, cols=1)
    body_tbl.style = 'Table Grid'
    bc = body_tbl.rows[0].cells[0]
    set_cell_bg(bc, "F8F9FA")
    bp = bc.paragraphs[0]
    bp.paragraph_format.space_before = Pt(6)
    bp.paragraph_format.space_after = Pt(6)
    bp.paragraph_format.left_indent = Pt(6)
    bp.paragraph_format.right_indent = Pt(6)
    br2 = bp.add_run(body_copy)
    br2.font.size = Pt(9)
    br2.font.color.rgb = DARK_GRAY

    # spacing
    sp = copy_cell.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)

    # CTA section
    cta_lp = copy_cell.add_paragraph()
    cta_lp.paragraph_format.left_indent = Pt(8)
    cta_lp.paragraph_format.space_before = Pt(2)
    cta_lp.paragraph_format.space_after = Pt(2)
    clr = cta_lp.add_run("CTA BUTTON")
    clr.font.bold = True
    clr.font.size = Pt(7.5)
    clr.font.color.rgb = GREEN_CTA

    cta_p = copy_cell.add_paragraph()
    cta_p.paragraph_format.left_indent = Pt(8)
    cta_p.paragraph_format.space_before = Pt(1)
    cta_p.paragraph_format.space_after = Pt(2)
    cta_r = cta_p.add_run(f"[ {cta_text} ]")
    cta_r.font.bold = True
    cta_r.font.size = Pt(11)
    cta_r.font.color.rgb = GREEN_CTA

    cta_sub_p = copy_cell.add_paragraph()
    cta_sub_p.paragraph_format.left_indent = Pt(8)
    cta_sub_p.paragraph_format.space_before = Pt(0)
    cta_sub_p.paragraph_format.space_after = Pt(8)
    cta_sub_r = cta_sub_p.add_run(cta_subtext)
    cta_sub_r.font.size = Pt(8)
    cta_sub_r.font.italic = True
    cta_sub_r.font.color.rgb = MID_GRAY

    # Platform notes
    plat_lp = copy_cell.add_paragraph()
    plat_lp.paragraph_format.left_indent = Pt(8)
    plat_lp.paragraph_format.space_before = Pt(0)
    plat_lp.paragraph_format.space_after = Pt(2)
    plr = plat_lp.add_run("PLACEMENT")
    plr.font.bold = True
    plr.font.size = Pt(7.5)
    plr.font.color.rgb = MID_GRAY

    plat_val = copy_cell.add_paragraph()
    plat_val.paragraph_format.left_indent = Pt(8)
    plat_val.paragraph_format.space_before = Pt(0)
    plat_val.paragraph_format.space_after = Pt(10)
    plvr = plat_val.add_run(platform_notes)
    plvr.font.size = Pt(8.5)
    plvr.font.color.rgb = DARK_GRAY

    # Remove table borders for outer table
    for row in tbl2.rows:
        for c in row.cells:
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '2')
                el.set(qn('w:color'), 'DDDDDD')
                tcBorders.append(el)
            tcPr.append(tcBorders)

    # Video script notes section
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    vn_tbl = doc.add_table(rows=1, cols=1)
    vn_tbl.style = 'Table Grid'
    vn_cell = vn_tbl.rows[0].cells[0]
    set_cell_bg(vn_cell, "EEF4FF")

    vn_p = vn_cell.add_paragraph()
    vn_p.paragraph_format.space_before = Pt(6)
    vn_p.paragraph_format.space_after = Pt(2)
    vn_p.paragraph_format.left_indent = Pt(8)
    vnl = vn_p.add_run("🎬  VIDEO SCRIPT NOTES")
    vnl.font.bold = True
    vnl.font.size = Pt(8)
    vnl.font.color.rgb = BRAND_BLUE

    vn_body = vn_cell.add_paragraph()
    vn_body.paragraph_format.space_before = Pt(2)
    vn_body.paragraph_format.space_after = Pt(8)
    vn_body.paragraph_format.left_indent = Pt(8)
    vn_body.paragraph_format.right_indent = Pt(8)
    vnbr = vn_body.add_run(video_script_notes)
    vnbr.font.size = Pt(9)
    vnbr.font.color.rgb = DARK_GRAY

    # remove table borders
    for row in vn_tbl.rows:
        for c in row.cells:
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:color'), 'C5D8FF')
                tcBorders.append(el)
            tcPr.append(tcBorders)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_divider(doc, "DDDDDD")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Cover / Title Block ───────────────────────────────────────────────────────
cover_tbl = doc.add_table(rows=1, cols=1)
cover_tbl.style = 'Table Grid'
cc = cover_tbl.rows[0].cells[0]
set_cell_bg(cc, "12141A")

cp1 = cc.add_paragraph()
cp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp1.paragraph_format.space_before = Pt(22)
cp1.paragraph_format.space_after = Pt(4)
r = cp1.add_run("THE EXCHANGE APP")
r.font.bold = True
r.font.size = Pt(22)
r.font.color.rgb = BRAND_GOLD

cp2 = cc.add_paragraph()
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp2.paragraph_format.space_before = Pt(0)
cp2.paragraph_format.space_after = Pt(6)
r2 = cp2.add_run("Facebook Video Ads — Creative Brief & Copy Journal")
r2.font.bold = True
r2.font.size = Pt(13)
r2.font.color.rgb = WHITE

cp3 = cc.add_paragraph()
cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp3.paragraph_format.space_before = Pt(0)
cp3.paragraph_format.space_after = Pt(6)
r3 = cp3.add_run("7 Ads — Website Builder Lead · Church Network · All-In-One Platform · 14-Day Trial · Retargeting")
r3.font.size = Pt(10)
r3.font.italic = True
r3.font.color.rgb = MID_GRAY

cp4 = cc.add_paragraph()
cp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp4.paragraph_format.space_before = Pt(0)
cp4.paragraph_format.space_after = Pt(16)
r4 = cp4.add_run("theexchangeapp.church   ·   No payroll focus   ·   March 2026")
r4.font.size = Pt(8.5)
r4.font.color.rgb = MID_GRAY

for row in cover_tbl.rows:
    for c in row.cells:
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'nil')
            tcBorders.append(el)
        tcPr.append(tcBorders)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ── How to use this document ──────────────────────────────────────────────────
add_paragraph_with_border(doc, "HOW TO USE THIS DOCUMENT", font_size=9, bold=True,
                           color=BRAND_BLUE, space_before=2, space_after=4)
notes = [
    "Each ad section below represents one Facebook video ad unit — ready to hand off to a video editor or AI video tool.",
    "The PAIN POINT HOOK is the opening line spoken on camera or shown as text in the first 3 seconds of the video.",
    "The AD BODY COPY is pasted directly into Facebook Ads Manager's 'Primary Text' field.",
    "The VIDEO SCRIPT NOTES describe what should appear on screen and when — use these with AI video tools (HeyGen, Runway, etc.).",
    "Start with AD 01 and AD 02 (the primary launch tests). Run these first. Pause the others until you have data.",
    "Images shown are Pexels reference images — use similar visuals or B-roll in the actual video ads.",
]
for n in notes:
    np = doc.add_paragraph(style='List Bullet')
    np.paragraph_format.space_before = Pt(1)
    np.paragraph_format.space_after = Pt(2)
    nr = np.runs[0] if np.runs else np.add_run(n)
    nr.font.size = Pt(8.5)
    nr.font.color.rgb = DARK_GRAY
    if not np.runs:
        np.add_run(n).font.size = Pt(8.5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_divider(doc, "1A73E8")
doc.add_paragraph().paragraph_format.space_after = Pt(14)

# ════════════════════════════════════════════════════════════════════════════
# AD 01 — PRIMARY: WEBSITE BUILDER — NO CHURCH WEBSITE
# ════════════════════════════════════════════════════════════════════════════
add_ad_section(
    doc=doc,
    ad_num=1,
    title="Your Church Has Been Putting Off a Real Website for Years. It Takes 10 Minutes.",
    pain_hook=(
        '"Have you ever sent someone your church\'s website\n'
        'and immediately felt embarrassed\n'
        'by what they were about to see?"'
    ),
    body_copy=(
        "Most churches have one of three website problems:\n\n"
        "→  No website at all\n"
        "→  An outdated one nobody knows how to update\n"
        "→  A Squarespace site that doesn't connect to anything else your church uses\n\n"
        "Your church's website is the first thing someone sees when they're looking for a home. "
        "It's the page a curious visitor Googles on Sunday morning before they decide whether to show up. "
        "It's where someone halfway across the country looks when they want to support your mission.\n\n"
        "And if it looks abandoned — or doesn't exist — that moment is gone.\n\n"
        "The Exchange gives your church a full website — with your own domain, "
        "a donation page built in, event listings, ministries pages, and a media archive — "
        "built in about 10 minutes, no coding required.\n\n"
        "Then everything connects: your website, your giving, your member management, "
        "your budget — all in one platform. One login.\n\n"
        "Free 14-day trial. No credit card. theexchangeapp.church"
    ),
    cta_text="Build Your Church Website Free",
    cta_subtext="14-day trial · No credit card · Your own domain · Done in 10 minutes",
    platform_notes=(
        "Facebook Feed + Instagram Feed + Instagram Reels\n"
        "Target: Pastors, Executive Pastors, Church Administrators (job title targeting)\n"
        "PRIMARY test ad — run this first at $333/mo"
    ),
    video_script_notes=(
        "0–3s:   HOOK text on dark background:\n"
        "        'Have you ever sent someone your church website and immediately regretted it?'\n\n"
        "3–8s:   Show a phone screen with a Google search for a church name. "
        "The result: a bare, outdated-looking page with a 2012 design.\n"
        "        Narrator: 'For a lot of churches, the website is the last thing anyone touches.'\n\n"
        "8–20s:  Screen recording: Open The Exchange website builder. "
        "Show drag-and-drop blocks being placed. Enter a church name. Upload a logo. "
        "Add a Give button. Select a domain. Done.\n"
        "        Narrator: 'The Exchange builds your entire church website in about 10 minutes. "
        "Your own domain. A donation page built in. Events, ministries, media — all of it.'\n\n"
        "20–27s: Pull back to show the full dashboard: Website tab, Giving tab, Members tab, Budget tab.\n"
        "        Narrator: 'And it connects to everything else your church needs — giving, members, "
        "budget — one platform, one login.'\n\n"
        "27–30s: CTA card: 'Free 14-Day Trial · Build Your Website Today · theexchangeapp.church'\n\n"
        "Captions required. 4:5 ratio for Feed. 9:16 for Reels. Hook in first 3 seconds is everything."
    ),
    image_path=os.path.join(IMAGES_DIR, "congregation.jpg"),
    image_caption="Reference: Church congregation — Pexels #34504326\n(Open with church interior, then cut to the website builder)",
    bg_hex_title="1A2340",
    phase="PHASE 1 — PRIMARY TEST AD  ★  RUN THIS FIRST"
)

# ════════════════════════════════════════════════════════════════════════════
# AD 02 — WEBSITE + CONNECTION: NEW MEMBERS CAN'T FIND YOU
# ════════════════════════════════════════════════════════════════════════════
add_ad_section(
    doc=doc,
    ad_num=2,
    title="If Someone New Moved to Your City This Weekend, Could They Find Your Church Online?",
    pain_hook=(
        '"Someone just moved to your city.\n'
        'They searched for a church on Sunday morning.\n'
        'They found your competitor — not you."'
    ),
    body_copy=(
        "Every week, people move to new cities and look for a church to call home.\n\n"
        "They Google. They check websites. They look for a service time, a location, "
        "a sense of the community. If your website is outdated, buried, "
        "or nonexistent — they never find you.\n\n"
        "But here's what most church platforms miss: "
        "it's not just about having a website.\n\n"
        "It's about being discoverable inside a network of churches.\n\n"
        "On The Exchange, your church has:\n\n"
        "✅  A full website — built in 10 minutes, your own domain\n"
        "✅  A profile in the Exchange church network — searchable by city and category\n"
        "✅  A digital connect form — new visitors join your church community from their phone\n"
        "✅  A community feed — your congregation stays connected between Sundays\n"
        "✅  Giving built in — from the same platform, same login\n\n"
        "Your church shouldn't be invisible. "
        "The Exchange makes sure it isn't.\n\n"
        "Free 14-day trial. No credit card. theexchangeapp.church"
    ),
    cta_text="Get Your Church Found — Start Free",
    cta_subtext="Website · Church directory · Connect forms · All in one platform",
    platform_notes=(
        "Facebook Feed + Instagram\n"
        "Target: Lead Pastors, Church Planters, Ministry Directors, Church Growth interest audience\n"
        "Strong for: churches in growing cities, new church plants, churches looking to grow"
    ),
    video_script_notes=(
        "0–3s:   HOOK on screen: 'If someone moved to your city this weekend — "
        "would they find your church online?'\n\n"
        "3–9s:   Show a person on a phone, Googling 'churches near me.' "
        "Results show several churches — one has a great website with photos, service times, "
        "and a 'Plan a Visit' button. The others are outdated or missing.\n"
        "        Narrator: 'Every week, people are looking for a church. "
        "The ones with a digital presence — they get found. The others don't.'\n\n"
        "9–20s:  Screen recording: The Exchange church profile page — "
        "shows the website, the connect form, the community feed.\n"
        "        Show a new visitor filling out a connect form on their phone.\n"
        "        Narrator: 'The Exchange gives your church a full website, "
        "a profile in the Exchange church network, and a digital connect form — "
        "so new visitors find you and join your community before they even walk through the door.'\n\n"
        "20–27s: Pull back to the full platform dashboard — website, giving, members, budget.\n"
        "        Narrator: 'Everything connects. One platform. Free to start.'\n\n"
        "27–30s: CTA: 'Get Your Church Found · Free 14-Day Trial · theexchangeapp.church'\n\n"
        "Tone: Warm urgency. Every week without this costs you real people."
    ),
    image_path=os.path.join(IMAGES_DIR, "church_cross.jpg"),
    image_caption="Reference: Church exterior — Pexels #33676861\n(Use exterior + phone search visual in video)",
    bg_hex_title="0F3D20",
    phase="PHASE 1 — PRIMARY TEST AD"
)

# ════════════════════════════════════════════════════════════════════════════
# AD 03 — ALL-IN-ONE: WEBSITE + EVERYTHING, ONE PLATFORM
# ════════════════════════════════════════════════════════════════════════════
add_ad_section(
    doc=doc,
    ad_num=3,
    title="Your Church Website, Your Giving, Your Members, and Your Budget — Why Are They in Four Different Places?",
    pain_hook=(
        '"Your church\'s website is on Squarespace.\n'
        'Your giving is on a separate platform.\n'
        'Your member list is in someone\'s phone.\n'
        'Your budget is in a spreadsheet.\n'
        'They don\'t talk to each other."'
    ),
    body_copy=(
        "This is the reality for most churches:\n\n"
        "→  Website on Squarespace or Wix — $20–$40/mo\n"
        "→  Online giving on a separate platform — $29–$500+/mo\n"
        "→  Member list scattered across texts, emails, and a spreadsheet\n"
        "→  Budget in a Google Sheet only the treasurer can access\n"
        "→  Announcements sent from a personal email or a third app\n\n"
        "Five tools. Five bills. Five logins. Nothing connected.\n"
        "And when something falls through the cracks — it usually does.\n\n"
        "The Exchange replaces all of it:\n\n"
        "✅  Full church website — drag-and-drop, your own domain, done in 10 minutes\n"
        "✅  Online giving — donation forms, recurring gifts, giving history, receipts\n"
        "✅  Member management — connect forms, activity tracking, follow-up\n"
        "✅  Shared budget sheet — your whole team sees the same numbers\n"
        "✅  Community feed — post updates, events, and campaigns all in one place\n\n"
        "One platform. One login. Starting at $0.\n\n"
        "Free 14-day trial. No credit card. theexchangeapp.church"
    ),
    cta_text="Try Everything Free for 14 Days",
    cta_subtext="All features unlocked · No credit card · One login for your whole team",
    platform_notes=(
        "Facebook Feed\n"
        "Target: Church Administrators, Office Managers, Associate Pastors\n"
        "People managing church operations day-to-day — they feel this pain most"
    ),
    video_script_notes=(
        "0–3s:   HOOK — text appears one line at a time on dark background:\n"
        "        'Website: Squarespace.'\n"
        "        'Giving: Tithe.ly.'\n"
        "        'Members: a spreadsheet.'\n"
        "        'Budget: another spreadsheet.'\n"
        "        'One church. Four tools. None connected.'\n\n"
        "3–12s:  Fast-cut montage: browser tabs switching — Squarespace, giving platform, "
        "Google Sheets, email client. Each tab has a price label.\n"
        "        Narrator: 'Most church staff are managing five tools that don't talk to each other. "
        "Every week, something falls through the cracks.'\n\n"
        "12–22s: Wipe to The Exchange. Show the left navigation:\n"
        "        Website → Giving → Members → Budget → Community Feed.\n"
        "        Fast click-through of each section.\n"
        "        Narrator: 'The Exchange is one platform for everything your church runs. "
        "Your website is connected to your giving, your members, your budget, and your community — "
        "all in one place, one login.'\n\n"
        "22–27s: Show the Free plan and $29 Growth plan side by side.\n"
        "        Narrator: 'Start free. 14-day trial of all premium features.'\n\n"
        "27–30s: CTA: 'One Platform · 14-Day Trial · theexchangeapp.church'\n\n"
        "Energy: 'This should be obvious. Why is everyone still doing it the hard way?'"
    ),
    image_path=os.path.join(IMAGES_DIR, "stressed_admin.jpg"),
    image_caption="Reference: Overwhelmed admin — Pexels #3808818\n(Open on scattered-tools chaos, resolve to the Exchange dashboard)",
    bg_hex_title="3D1A0A",
    phase="PHASE 1 — ALL-IN-ONE PLATFORM AD"
)

# ════════════════════════════════════════════════════════════════════════════
# AD 04 — CONNECTING CHURCHES: THE PLATFORM NETWORK
# ════════════════════════════════════════════════════════════════════════════
add_ad_section(
    doc=doc,
    ad_num=4,
    title="What If Your Church Was Part of a Network Where Churches Actually Supported Each Other?",
    pain_hook=(
        '"Your church is doing meaningful work.\n'
        'But outside your four walls,\n'
        'nobody in the broader church community\n'
        'can find you, follow you, or support you."'
    ),
    body_copy=(
        "The Church wasn't designed to operate in isolation.\n\n"
        "But right now, most churches exist as individual islands — "
        "each with their own website, their own giving platform, "
        "their own member database — "
        "with no shared infrastructure connecting them.\n\n"
        "The Exchange is building that infrastructure.\n\n"
        "On The Exchange, your church isn't just a website. "
        "It's a profile inside a network of churches, ministries, and nonprofits "
        "that can discover each other, connect, post updates, support campaigns, "
        "and build the kind of community the Body of Christ was always meant to have.\n\n"
        "→  Your church profile is searchable inside the Exchange network\n"
        "→  Other churches and members can find your mission and support it directly\n"
        "→  Post updates to your community feed — members stay engaged between Sundays\n"
        "→  Connect forms bring new people into your church digitally\n"
        "→  Your website, giving, members, and budget all live in the same place\n\n"
        "This isn't just a website builder. "
        "It's the operating system for churches that want to grow together.\n\n"
        "Free 14-day trial. No credit card. theexchangeapp.church"
    ),
    cta_text="Join the Exchange Network — Free",
    cta_subtext="Church website · Network profile · Community feed · All in one platform",
    platform_notes=(
        "Facebook Feed + Instagram\n"
        "Target: Lead Pastors, Church Planters, Denominational network leaders\n"
        "Interests: Church community, Ministry networking, Church growth, Faith collaboration"
    ),
    video_script_notes=(
        "0–3s:   HOOK on screen: 'What if your church was part of a network — "
        "where other churches could actually find and support you?'\n\n"
        "3–10s:  B-roll montage: multiple different churches — a small congregation, "
        "a suburban church, a church plant meeting in a coffee shop. Each shown briefly, "
        "each isolated. No connection between them visually.\n"
        "        Narrator: 'Most churches operate alone. "
        "Each one with their own website, their own tools, their own digital island.'\n\n"
        "10–22s: Screen recording: The Exchange Discover page.\n"
        "        Show churches listed — photos, descriptions, locations, Give buttons.\n"
        "        Show a community feed: a church posts an update, others like and comment.\n"
        "        Show a connect form being filled out by a new visitor on their phone.\n"
        "        Narrator: 'The Exchange connects your church to a network. "
        "Other churches can find you. New visitors can connect digitally. "
        "Your community stays active between Sundays — all on one platform with your website.'\n\n"
        "22–27s: Show the full dashboard — Website, Giving, Members, Budget, Community.\n"
        "        Narrator: 'Your website. Your community. Your giving. One place.'\n\n"
        "27–30s: CTA: 'Free 14-Day Trial · theexchangeapp.church'\n\n"
        "Tone: Vision + practical. 'This is what it should have always looked like.'"
    ),
    image_path=os.path.join(IMAGES_DIR, "volunteers.jpg"),
    image_caption="Reference: Community/volunteers — Pexels #7156193\n(Use churches-connecting imagery in video)",
    bg_hex_title="2A0A3D",
    phase="PHASE 1 — NETWORK + COMMUNITY AD"
)

# ════════════════════════════════════════════════════════════════════════════
# AD 05 — MEMBER ENGAGEMENT: CONGREGATION DISCONNECTED BETWEEN SUNDAYS
# ════════════════════════════════════════════════════════════════════════════
add_ad_section(
    doc=doc,
    ad_num=5,
    title="Your Congregation Gathers on Sunday. Where Do They Go the Other Six Days?",
    pain_hook=(
        '"Does your church have a way to stay\n'
        'connected with your members on Monday,\n'
        'Tuesday, Wednesday — or do they just\n'
        'disappear until next Sunday?"'
    ),
    body_copy=(
        "Sunday is one day. Your congregation lives the other six.\n\n"
        "Most churches have no digital presence that keeps members engaged "
        "between services. No feed to post to. No way for members to connect with each other. "
        "No tool for a new visitor to say 'I want to be part of this community' "
        "without showing up in person and hoping someone notices.\n\n"
        "The Exchange changes that:\n\n"
        "→  Your church website — the front door, always open, any device\n"
        "→  Connect forms — new visitors join your church digitally from their phone\n"
        "→  Community feed — post updates, events, and announcements your members actually see\n"
        "→  AI surveys from your sermon notes — members engage with the message mid-week\n"
        "→  Member activity tracking — know who's engaged and who needs a follow-up\n\n"
        "Church doesn't end on Sunday. Your platform shouldn't either.\n\n"
        "Free 14-day trial. No credit card. theexchangeapp.church"
    ),
    cta_text="Keep Your Church Connected — Start Free",
    cta_subtext="Website · Connect forms · Community feed · Member management · One platform",
    platform_notes=(
        "Facebook Feed + Instagram\n"
        "Target: Lead Pastors, Campus Pastors, Associate Pastors focused on community/growth\n"
        "Interests: Church community, Congregation engagement, Discipleship, Church growth"
    ),
    video_script_notes=(
        "0–3s:   HOOK: 'Your congregation gathers on Sunday. "
        "What happens to your community the other six days?'\n\n"
        "3–9s:   B-roll: A Sunday morning church service — packed, energetic, engaged. "
        "Then: cut to Monday morning. A church building, quiet. Locked. Empty parking lot.\n"
        "        Narrator: 'For most churches, the community only exists on Sunday morning. "
        "There's no digital space that keeps people connected the rest of the week.'\n\n"
        "9–20s:  Screen recording: Exchange platform.\n"
        "        Show the community feed — a pastor posting mid-week encouragement. "
        "Show a connect form being filled out on a phone by someone who visited Sunday. "
        "Show AI survey questions generated from a sermon being sent to members.\n"
        "        Narrator: 'The Exchange gives your church a digital home between Sundays. "
        "A community feed. Digital connect forms. Sermon-based surveys. "
        "And your website — the always-open front door — all in one platform.'\n\n"
        "20–27s: Show member activity dashboard — who's engaged, who's been inactive.\n"
        "        Narrator: 'And you can see exactly who's active and who needs a call.'\n\n"
        "27–30s: CTA: 'Free 14-Day Trial · theexchangeapp.church'\n\n"
        "Tone: Pastoral. This is about your people, not the software."
    ),
    image_path=os.path.join(IMAGES_DIR, "pastor.jpg"),
    image_caption="Reference: Pastor/congregation — Pexels #8815004\n(B-roll: Sunday service full, then cut to empty Monday building)",
    bg_hex_title="1A3D2A",
    phase="PHASE 1 — COMMUNITY ENGAGEMENT AD"
)

# ════════════════════════════════════════════════════════════════════════════
# AD 06 — TESTIMONIAL TEMPLATE (Record after pilot church onboarded)
# ════════════════════════════════════════════════════════════════════════════
add_ad_section(
    doc=doc,
    ad_num=6,
    title="Testimonial Ad — Pilot Church Story (Shoot After Pilot Phase, Month 1–2)",
    pain_hook=(
        '"We had a website on one platform,\n'
        'giving on another, and no way to connect\n'
        'it all. The Exchange put everything\n'
        'in one place in one afternoon."\n'
        '— [Pastor Name], [Church Name]'
    ),
    body_copy=(
        "[CHURCH NAME] in [CITY] was managing their church across four different platforms.\n\n"
        "Their website was on Squarespace. Their giving was separate. "
        "Their member list lived in a spreadsheet. "
        "New visitors had no digital way to connect.\n\n"
        "In one afternoon, they moved everything to The Exchange:\n\n"
        "→  Their church website is live — their own domain, donation page included\n"
        "→  New visitors fill out a connect form from their phone after Sunday service\n"
        "→  Their congregation stays engaged through a community feed all week\n"
        "→  Online giving is active — donors give from any device, any time\n"
        "→  The whole team manages members and budget from one dashboard\n\n"
        "Their pastor set it up himself. No technical background. No developer needed.\n\n"
        "Free 14-day trial. No credit card.\n"
        "→ theexchangeapp.church"
    ),
    cta_text="Start Your Free Trial",
    cta_subtext="See what your church looks like on one platform — in one afternoon",
    platform_notes=(
        "Facebook Feed + YouTube Pre-Roll\n"
        "Target: All cold audiences + retargeting warm audiences\n"
        "Testimonial ads consistently outperform all other formats — prioritize after pilot church is onboarded"
    ),
    video_script_notes=(
        "FORMAT: Talking-head testimonial. Unscripted. Pastor or admin on camera.\n\n"
        "RECORD IN A SINGLE SESSION — ask these questions:\n"
        "1. 'What tools were you using before, and what was the biggest headache?'\n"
        "2. 'How long did it take you to get the website up and everything set up?'\n"
        "3. 'What's the feature that surprised you the most?'\n"
        "4. 'How has your congregation responded to the digital connect form or community feed?'\n"
        "5. 'What would you say to another pastor who's still on five different platforms?'\n\n"
        "EDIT STRUCTURE (30–45 seconds):\n"
        "0–5s:   B-roll of church exterior or Sunday service in action\n"
        "5–15s:  Pastor on camera answering question 1 — the before (the pain)\n"
        "15–25s: Screen recording of their actual Exchange website + community feed + giving page\n"
        "25–35s: Pastor on camera answering question 5 — the recommendation\n"
        "35–40s: CTA card: 'Free 14-Day Trial · theexchangeapp.church'\n\n"
        "DO NOT SCRIPT. Raw and honest always converts better than polished."
    ),
    image_path=os.path.join(IMAGES_DIR, "pastor.jpg"),
    image_caption="Reference: Pastor preaching — Pexels #8815004\n(Replace with real pilot church footage when available)",
    bg_hex_title="0A1A3D",
    phase="PHASE 2 — TESTIMONIAL (After Pilot Church Onboarded)"
)

# ════════════════════════════════════════════════════════════════════════════
# AD 07 — RETARGETING: TRIAL ENDING — DON'T LOSE YOUR WEBSITE
# ════════════════════════════════════════════════════════════════════════════
add_ad_section(
    doc=doc,
    ad_num=7,
    title="Retargeting — Your Trial Ends in 2 Days. Your Church Website Goes Dark If You Don't Upgrade.",
    pain_hook=(
        '"You built your church website during your trial.\n'
        'In 2 days, it goes offline\n'
        'unless you upgrade.\n'
        'Your domain. Your giving page. All of it."'
    ),
    body_copy=(
        "Your 14-day trial is almost over.\n\n"
        "During your trial, you built something real:\n"
        "→  Your church website — live at your custom domain\n"
        "→  Your giving page — already shared with your congregation\n"
        "→  Your connect form — new visitors are already filling it out\n"
        "→  Your community feed — members are already following along\n\n"
        "When the trial ends, your website goes offline. "
        "Your giving page goes dark. The connect form stops working.\n\n"
        "The Growth plan keeps everything live. It's $29/mo.\n"
        "That's one dollar a day to keep your church's entire digital presence running.\n\n"
        "Your website. Your giving. Your community. Your members.\n"
        "Everything you built — stays exactly as you left it.\n\n"
        "Upgrade before your trial ends.\n"
        "Month-to-month. Cancel any time."
    ),
    cta_text="Keep My Church Website Live — $29/mo",
    cta_subtext="Everything you built stays live · Month-to-month · Cancel any time",
    platform_notes=(
        "Facebook Feed + Instagram\n"
        "Audience: Trial users on Day 12–13 who have NOT upgraded\n"
        "Highest-intent audience — maximize daily budget here during trial window"
    ),
    video_script_notes=(
        "FORMAT: Calm urgency. 20–25 seconds max. They know the product.\n\n"
        "0–3s:   Text on screen, large and bold: '2 days left on your trial.'\n\n"
        "3–10s:  Screen recording: the church website they built — "
        "showing the home page, the Give button, the events page.\n"
        "        Narrator: 'Your church website is live. Your giving page is active. "
        "Your community can already see it.'\n\n"
        "10–18s: Show the pricing page. Cursor moves to the $29 Growth plan.\n"
        "        Narrator: 'Keep all of it live for $1 a day. "
        "Month-to-month. Cancel any time. Nothing disappears.'\n\n"
        "18–25s: CTA card: 'Keep Your Church Website Live · $29/mo · theexchangeapp.church'\n\n"
        "Tone: Matter-of-fact. Not fear — just the practical reality. "
        "'You did the work. $1/day keeps it running.'"
    ),
    image_path=os.path.join(IMAGES_DIR, "laptop_work.jpg"),
    image_caption="Reference: Laptop/work — Pexels #7309483\n(Show Exchange website they built on the screen in video)",
    bg_hex_title="3D2A0A",
    phase="PHASE 2 — RETARGETING (Trial Day 12–13, Highest Intent)"
)

# ── A/B Test Schedule Summary ─────────────────────────────────────────────────
doc.add_page_break()

add_paragraph_with_border(doc, "A/B TEST SCHEDULE — RUN ONE VARIABLE AT A TIME",
                           font_size=11, bold=True, color=BRAND_BLUE, space_before=8, space_after=8)

test_data = [
    ("TEST 1\nMonth 2–3\nHook Copy",
     "AD 01: 'Embarrassed by your church website?'\nAD 02: 'Can a new visitor find you online?'\nAD 03: 'Five tools, none connected'",
     "$333/mo each\n($1,000 total)",
     "Pause the two lowest CPL performers after 2 weeks.\nDouble budget on the winner."),
    ("TEST 2\nMonth 3\nCreative Format",
     "A: Screen recording — website builder being built\nB: Talking-head testimonial (pilot church pastor)\nC: Tab-switching montage — five tools → one platform",
     "Split the $1,000\nequally across 3",
     "Winner = lowest cost per trial start.\nKill losers after 7 days if CTR < 1%."),
    ("TEST 3\nMonth 3–4\nAudience",
     "A: Job titles — Pastor, Church Admin, Ministry Director\nB: Interests — Church management, Church growth, Ministry\nC: Broad 25–65 US (let Meta find the audience)",
     "Same budget,\nwinning creative",
     "Run winner creative against all 3 audiences.\nPause the two worst after 3 weeks."),
    ("TEST 4\nMonth 4\nCTA",
     "A: 'Build Your Church Website Free'\nB: 'Try Everything Free for 14 Days'\nC: 'Get Your Church Found — Start Free'",
     "Same budget,\nwinning audience",
     "Lock in the winning CTA for all scale spend.\nWebsite builder CTA historically outperforms generic 'Start Free' for SaaS."),
]

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
headers = ["TEST PHASE", "WHAT WE'RE TESTING", "BUDGET SPLIT", "DECISION RULE"]
for i, h in enumerate(headers):
    c = tbl.rows[0].cells[i]
    set_cell_bg(c, "1A2340")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(h)
    r.font.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = WHITE

for row_data in test_data:
    row = tbl.add_row()
    for i, cell_text in enumerate(row_data):
        c = row.cells[i]
        set_cell_bg(c, "F8F9FA" if i % 2 == 0 else "FFFFFF")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(4)
        r = p.add_run(cell_text)
        r.font.size = Pt(8.5)
        r.font.color.rgb = DARK_GRAY

doc.add_paragraph().paragraph_format.space_after = Pt(16)

# ── Performance Benchmarks ────────────────────────────────────────────────────
add_paragraph_with_border(doc, "PERFORMANCE BENCHMARKS — WHEN TO SCALE vs. PAUSE",
                           font_size=11, bold=True, color=BRAND_BLUE, space_before=8, space_after=8)

bench_tbl = doc.add_table(rows=1, cols=4)
bench_tbl.style = 'Table Grid'
bench_headers = ["METRIC", "PAUSE IF BELOW / ABOVE", "TARGET", "NOTES"]
for i, h in enumerate(bench_headers):
    c = bench_tbl.rows[0].cells[i]
    set_cell_bg(c, "0F3D20")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(h)
    r.font.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = WHITE

bench_rows = [
    ("Click-Through Rate (CTR)",  "Below 1.0%",  "2.0%+",    "Below 1% = ad isn't resonating. Kill it."),
    ("Cost Per Click (CPC)",      "Above $3.00", "Under $1.50", "Facebook CPM ~$14.90 → need strong CTR"),
    ("Cost Per Lead (CPL)",       "Above $80",   "Under $40",   "Lead = free account signup completed"),
    ("Trial Start Rate",          "Below 15%",   "30%+",        "% of signups who start a 14-day trial"),
    ("Trial-to-Paid Conversion",  "Below 15%",   "25%+",        "Primary revenue gate"),
    ("Landing Page Conv. Rate",   "Below 5%",    "10–15%",      "Track with Microsoft Clarity heatmaps"),
]
for row_data in bench_rows:
    row = bench_tbl.add_row()
    for i, cell_text in enumerate(row_data):
        c = row.cells[i]
        set_cell_bg(c, "F8F9FA" if i % 2 == 0 else "FFFFFF")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Pt(4)
        r = p.add_run(cell_text)
        r.font.size = Pt(8.5)
        r.font.color.rgb = DARK_GRAY
        if i == 1:
            r.font.color.rgb = PAIN_RED
            r.font.bold = True
        elif i == 2:
            r.font.color.rgb = GREEN_CTA
            r.font.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(20)

# ── Footer ────────────────────────────────────────────────────────────────────
add_divider(doc, "1A73E8")
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(10)
fr = footer_p.add_run(
    "The Exchange App  ·  theexchangeapp.church\n"
    "Built on the theology of the Great Exchange — 2 Corinthians 5:21\n"
    '"God made him who had no sin to be sin for us, so that in him we might become the righteousness of God."'
)
fr.font.size = Pt(8.5)
fr.font.italic = True
fr.font.color.rgb = MID_GRAY

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"✅  Document saved: {OUTPUT_PATH}")
