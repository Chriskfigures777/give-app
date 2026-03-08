from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Colours ───────────────────────────────────────────────────────────────────
DARK       = RGBColor(0x0E, 0x11, 0x18)
EMERALD    = RGBColor(0x10, 0xB9, 0x81)
EMERALD_DK = RGBColor(0x06, 0x7A, 0x55)
SLATE      = RGBColor(0x47, 0x55, 0x69)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def ps(para, before=0, after=5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)

def add_h1(text):
    p   = doc.add_paragraph()
    ps(p, before=18, after=6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = EMERALD_DK
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "10B981")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_h2(text):
    p   = doc.add_paragraph()
    ps(p, before=12, after=3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK

def add_h3(text):
    p   = doc.add_paragraph()
    ps(p, before=8, after=3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = EMERALD

def add_body(text):
    p   = doc.add_paragraph()
    ps(p, before=0, after=5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = SLATE

def add_bullet(text):
    p   = doc.add_paragraph(style="List Bullet")
    ps(p, before=0, after=3)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = SLATE

def add_callout(text):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shade_cell(cell, "F0FDF4")
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   "single")
        tag.set(qn("w:sz"),    "8")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "10B981")
        tcBorders.append(tag)
    tcPr.append(tcBorders)
    cell.paragraphs[0].clear()
    p   = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = EMERALD_DK
    run.bold = True
    doc.add_paragraph()

def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, "0E1118")
        cell.paragraphs[0].clear()
        p   = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg  = "FFFFFF" if r_idx % 2 == 0 else "F9FAFB"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            shade_cell(cell, bg)
            cell.paragraphs[0].clear()
            p   = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = SLATE
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


# =============================================================================
# COVER PAGE
# =============================================================================
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
cover.paragraph_format.space_after  = Pt(0)
r = cover.add_run("Exchange Banking")
r.bold = True; r.font.size = Pt(32); r.font.color.rgb = DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(4)
p2.paragraph_format.space_after  = Pt(0)
r2 = p2.add_run("Confidential Growth Plan")
r2.font.size = Pt(18); r2.font.color.rgb = EMERALD

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(4)
p3.paragraph_format.space_after  = Pt(0)
r3 = p3.add_run("$100K / Month by Year Two")
r3.font.size = Pt(14); r3.font.color.rgb = SLATE

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(60)
r4 = p4.add_run("Prepared March 2026  \u00b7  Internal Use Only")
r4.font.size = Pt(9); r4.font.color.rgb = SLATE; r4.italic = True

doc.add_page_break()


# =============================================================================
# 1. CODEBASE ASSESSMENT
# =============================================================================
add_h1("1. What Is Already Built")
add_body(
    "A full technical review of the give-app-main codebase was conducted in March 2026. "
    "The platform is substantially more complete than most pre-seed companies at this stage. "
    "The table below reflects confirmed, functional modules."
)
add_table(
    ["Feature / Module", "Status"],
    [
        ["Donation platform (Stripe Connect, 1% fee)", "\u2705 Built"],
        ["Subscription billing \u2014 Free / Growth $29 / Pro $49 + 14-day trial", "\u2705 Built"],
        ["Website builder \u2014 templates, CMS, publish, custom domains", "\u2705 Built"],
        ["Social feed \u2014 posts, reactions, comments, donation sharing", "\u2705 Built"],
        ["Peer connections & real-time chat / messaging", "\u2705 Built"],
        ["Unit banking integration \u2014 customer creation, deposit accounts", "\u2705 Built (sandbox)"],
        ["Plaid bank account linking", "\u2705 Built"],
        ["Payment splits \u2014 missionaries, partners", "\u2705 Built"],
        ["Campaigns, goals, embeds, QR codes, donation links", "\u2705 Built"],
        ["Events + Eventbrite integration", "\u2705 Built"],
        ["Transactional email via Resend", "\u2705 Built"],
        ["Custom domain management \u2014 Route53, CloudFront, ACM", "\u2705 Built"],
        ["Admin panel", "\u2705 Built"],
        ["Survey builder (Typeform/Tally-style)", "\u2705 Built"],
        ["BankGO redirect / auth handoff", "\u2705 Built"],
        ["AI features \u2014 Anthropic SDK integrated", "\u2705 Built"],
        ["Dwolla + Astra payment webhooks", "\u2705 Built"],
        ["Advisor network / listing pages", "\u274c Not yet built"],
        ["Referral / Banking Leader program mechanics", "\u274c Not yet built"],
        ["Boost post paid feature", "\u274c Not yet built"],
        ["Mobile app \u2014 iOS / Android", "\u274c Not yet built"],
        ["BankGO banking UI (separate codebase)", "\u274c Not yet built"],
        ["Unit banking in production (currently sandbox)", "\u274c Pending compliance"],
    ],
    col_widths=[4.2, 2.1]
)
add_callout(
    "Key insight: The hardest technical work is already done. "
    "The remaining gaps are features (advisor network, referral program, boost) and compliance (Unit production). "
    "This is a product launch problem, not a build problem."
)


# =============================================================================
# 2. CAN THIS BE PROFITABLE
# =============================================================================
add_h1("2. Can This Be Profitable?")
add_body(
    "Yes \u2014 but the path is narrower and slower than the original plan projects. "
    "The business model is proven (Pushpay built to $150M/year targeting churches on a similar stack). "
    "Below is an honest stress-test of the two biggest revenue assumptions."
)
add_h2("SaaS Subscriptions \u2014 Stress Test")
add_table(
    ["Assumption", "Plan Says", "Realistic"],
    [
        ["Paying orgs at Month 24", "~900", "~500\u2013700 without raise"],
        ["Trial-to-paid conversion", "10\u201312%", "10\u201315% with good onboarding"],
        ["Trials needed per month at Month 18", "~200", "200\u2013300"],
        ["CAC per trial (Meta + Google)", "$50\u2013$80", "$100\u2013$150"],
        ["Ad spend needed to hit target", "$15\u201320K/mo", "$20\u201340K/mo"],
    ],
    col_widths=[3.0, 1.5, 1.8]
)
add_h2("Interchange Revenue \u2014 Stress Test")
add_table(
    ["Assumption", "Plan Says", "Realistic"],
    [
        ["Active debit users at Month 21", "3,000", "1,000\u20132,000"],
        ["Avg monthly spend per user", "$300", "$300 (defensible)"],
        ["Interchange rate", "1.2%", "1.0\u20131.3% (Unit dependent)"],
        ["Monthly interchange at 3,000 users", "$10,800", "$10,800"],
        ["Plan stated target", "$30,000/mo", "Requires 7,500+ active users"],
    ],
    col_widths=[3.0, 1.5, 1.8]
)
add_callout(
    "$40,000\u2013$65,000/month by Month 24 is realistic without outside capital. "
    "$80,000\u2013$110,000/month is achievable with a $300K\u2013$500K raise deployed into ads in Months 15\u201320. "
    "The $100K target is real \u2014 it just requires fuel."
)


# =============================================================================
# 3. PHASE 0
# =============================================================================
add_h1("3. Phase 0 \u2014 Production Ready (Months 1\u20132)")
add_h3("Budget: $3,000\u2013$6,000 one-time")
add_body(
    "You cannot run paid ads to a product that is not battle-tested. "
    "Phase 0 closes the gap between 'code exists' and 'product works reliably for strangers.' "
    "Every dollar spent on ads before this is partially wasted."
)
add_h2("What Must Happen in Phase 0")
add_bullet("Move Unit banking from sandbox to production \u2014 start the partner application immediately; approval takes 4\u20138 weeks.")
add_bullet("End-to-end QA the 14-day trial \u2192 billing \u2192 cancellation \u2192 win-back flow.")
add_bullet("QA the website builder publish + custom domain flow on real, non-dev orgs.")
add_bullet("Build and activate automated onboarding email sequences: Day 1, 3, 7, 12, 14.")
add_bullet("Deploy to production: Vercel + Supabase + AWS stack confirmed live.")
add_bullet("Instrument analytics: PostHog or Mixpanel so trial conversion is measurable from Day 1 of paid ads.")
add_h2("Phase 0 Cost Breakdown")
add_table(
    ["Item", "Cost"],
    [
        ["Vercel Pro + Supabase Pro + AWS (monthly, ongoing)", "$200\u2013$400/mo"],
        ["Email sequence copywriting + setup in Resend", "$500\u2013$1,500"],
        ["Analytics tooling (PostHog cloud)", "$0\u2013$150/mo"],
        ["Unit compliance consultant (strongly recommended)", "$3,000\u2013$5,000 one-time"],
        ["QA contractor (optional)", "$1,000\u2013$3,000"],
    ],
    col_widths=[3.8, 2.5]
)


# =============================================================================
# 4. PHASE 1
# =============================================================================
add_h1("4. Phase 1 \u2014 Website Builder Launch (Months 3\u20136)")
add_h3("Ad Budget: $5,000\u2013$8,000/month  \u00b7  Total: $20,000\u2013$30,000")
add_body(
    "The website builder is the acquisition anchor. The 14-day free trial converts because it delivers "
    "immediate, tangible value \u2014 a live church website within hours of sign-up. Paid advertising "
    "drives churches into the trial funnel; the trial converts them to Growth plan subscribers."
)
add_h2("Paid Ad Strategy \u2014 Month 3\u20136")
add_table(
    ["Channel", "Monthly Budget", "What to Run"],
    [
        ["Meta (Facebook / Instagram)", "$3,000\u2013$4,000", "Video ads: church website built in under 10 minutes. Target church admins 35\u201365, Christian interest groups, church pages. Lookalike from existing user list."],
        ["Google Search", "$1,500\u2013$2,500", "High-intent keywords: 'church website builder', 'church giving software', 'online church donations'. Lower volume, higher conversion."],
        ["YouTube pre-roll", "$500\u2013$1,000", "30-second product demo targeting faith/church channels."],
    ],
    col_widths=[1.8, 1.5, 3.0]
)
add_h2("Realistic Acquisition Math at $5K/Month")
add_table(
    ["Metric", "Value"],
    [
        ["Monthly ad spend", "$5,000"],
        ["CAC per trial (realistic)", "$100\u2013$150"],
        ["Trials per month", "33\u201350"],
        ["Trial-to-paid conversion rate", "12%"],
        ["New paying orgs per month", "4\u20136"],
        ["New MRR added per month", "$116\u2013$174"],
        ["Cumulative paying orgs by Month 6", "~25\u201335"],
        ["Subscription MRR at Month 6", "$725\u2013$1,015/mo"],
    ],
    col_widths=[3.2, 3.1]
)
add_h2("Organic Acquisition (Runs Alongside Ads \u2014 Free)")
add_bullet("'Powered by Exchange' footer on every published church site creates compounding SEO from Month 3.")
add_bullet("Ask first 10 paying customers to post their site URL to their church Facebook page.")
add_bullet("Faith-based niche podcast sponsorships: $500\u2013$1,500/episode, highly targeted.")
add_bullet("Founder direct outreach via LinkedIn to church administrators.")
add_bullet("Submit to church tech roundups and Christian business directories.")
add_h2("Month 6 Revenue Target \u2014 Realistic")
add_table(
    ["Stream", "Monthly Amount"],
    [
        ["~30 Growth orgs \u00d7 $29/mo", "$870"],
        ["1% fee on ~$50K donated through platform", "$500"],
        ["Total MRR", "~$1,370/mo"],
    ],
    col_widths=[3.8, 2.5]
)
add_callout(
    "Phase 1 is not about revenue \u2014 it is about finding what converts. "
    "Month 3\u20136 reveals your real CAC, your best ad creative, and which onboarding steps drive trial-to-paid. "
    "That data is worth more than the MRR."
)


# =============================================================================
# 5. PHASE 2
# =============================================================================
add_h1("5. Phase 2 \u2014 Mobile App + BankGO Launch (Months 7\u201312)")
add_h3("Dev Cost: $15,000\u2013$30,000  \u00b7  Ad Budget: $8,000\u2013$12,000/month")
add_h2("Mobile App \u2014 Fastest Path")
add_body(
    "Your existing Next.js stack makes the fastest path to the App Store a React Native WebView shell \u2014 "
    "a native app container that loads your existing web app. Gets you into iOS and Android in "
    "6\u201310 weeks rather than 6 months."
)
add_table(
    ["Item", "Timeline", "Cost"],
    [
        ["iOS + Android WebView shell", "6\u201310 weeks", "$8,000\u2013$15,000"],
        ["Native payment \u2014 Apple Pay, Google Pay", "+2\u20133 weeks", "$3,000\u2013$5,000"],
        ["Push notifications (React Native)", "Included", "$0"],
        ["App Store / Play Store review", "4\u20138 weeks", "Time only"],
    ],
    col_widths=[2.5, 1.5, 1.8]
)
add_h2("BankGO Production Launch \u2014 The Highest-Risk Item")
add_body(
    "Unit banking in production requires formal compliance work that cannot be shortcut. "
    "The single most important action in this entire plan is: start the Unit production application now."
)
add_table(
    ["Requirement", "Timeline / Cost"],
    [
        ["Formal Unit partner agreement", "Start now \u2014 takes 4\u201312 weeks for approval"],
        ["BSA/AML compliance policy documentation", "$2,000\u2013$5,000 with consultant"],
        ["KYC flow QA for every state", "2\u20133 weeks internal"],
        ["FDIC pass-through insurance documentation", "Included in Unit partner process"],
        ["Full compliance consulting (recommended)", "$5,000\u2013$10,000 one-time"],
    ],
    col_widths=[3.0, 3.3]
)
add_callout(
    "If you wait until Month 5 to apply for Unit production access, you will not launch banking by Month 12. "
    "Start the application during Phase 0 (Month 1)."
)
add_h2("Month 12 Revenue Target \u2014 Realistic")
add_table(
    ["Stream", "Monthly Amount"],
    [
        ["150 Growth orgs \u00d7 $29/mo", "$4,350"],
        ["20 Pro orgs \u00d7 $49/mo", "$980"],
        ["1% fees on $200K/mo platform donation volume", "$2,000"],
        ["200 BankGO debit users \u00d7 $300 avg spend \u00d7 1.2% interchange", "$720"],
        ["Total MRR", "~$8,050/mo"],
    ],
    col_widths=[3.8, 2.5]
)


# =============================================================================
# 6. PHASE 3
# =============================================================================
add_h1("6. Phase 3 \u2014 Advisor Network + Growth Scaling (Months 13\u201318)")
add_h3("Ad Budget: $12,000\u2013$18,000/month")
add_body(
    "By Month 13, your analytics will reveal which channels work and your real CAC. "
    "Scale spend behind proven funnels, add the advisor network, and activate the Banking Leader referral program."
)
add_h2("Advisor Network \u2014 What to Build")
add_body("Estimated 3\u20134 weeks of focused development on the existing codebase.")
add_bullet("Advisor application portal (form + review flow)")
add_bullet("Advisor profile pages with discovery map inside BankGO")
add_bullet("Advisor billing at $99/mo founding cohort \u2192 $120/mo standard")
add_bullet("Analytics for advisors: profile views, clicks, inquiry counts")
add_h2("Do NOT Launch the Advisor Network Before 100 Paying Orgs")
add_body(
    "Advisors pay $99\u2013$150/month because they are reaching a verified, faith-aligned, "
    "financially-engaged audience. If the platform has 30 orgs, advisors will not convert \u2014 "
    "and early rejection damages the feature's long-term credibility. Wait until Month 13 minimum."
)
add_h2("Banking Leader Referral Program")
add_body(
    "The referral dashboard, unique code system, and leader-tier mechanics are not yet in the codebase. "
    "This is 2\u20133 weeks of development. Launch alongside the advisor network in Month 13."
)
add_h2("Month 18 Revenue Target \u2014 Realistic")
add_table(
    ["Stream", "Monthly Amount"],
    [
        ["400 paying orgs blended ~$32 ARPU", "$12,800"],
        ["1% fees on $500K/mo donation volume", "$5,000"],
        ["800 BankGO users \u00d7 $300 spend \u00d7 1.2% interchange", "$2,880"],
        ["25 advisors \u00d7 $99/mo", "$2,475"],
        ["Boost campaigns (estimated)", "$1,500"],
        ["Total MRR", "~$24,655/mo"],
    ],
    col_widths=[3.8, 2.5]
)


# =============================================================================
# 7. PHASE 4
# =============================================================================
add_h1("7. Phase 4 \u2014 $100K Run Rate (Months 19\u201324)")
add_body(
    "Phase 4 is not a single launch \u2014 it is the compounding of Phases 1\u20133. "
    "By Month 19, all five revenue streams are active. The focus shifts from building to scaling."
)
add_h2("To Hit $100K/Month by Month 24 You Need One Of:")
add_bullet("A referral flywheel generating 30+ organic org sign-ups/month")
add_bullet("BankGO debit user base reaching 2,000\u20133,000+ active users")
add_bullet("A $300K\u2013$500K raise deployed into ads at $20\u201340K/month in Months 15\u201320")
add_h2("Month 24 Revenue Target \u2014 Revised Realistic")
add_table(
    ["Stream", "Self-Funded", "With $300K Raise"],
    [
        ["Paying orgs \u00d7 ARPU", "$29,700", "$42,000"],
        ["1% donation fees", "$15,000", "$22,000"],
        ["BankGO interchange", "$7,200", "$14,400"],
        ["Advisor subscriptions", "$6,000", "$9,600"],
        ["Boost campaigns", "$3,500", "$5,000"],
        ["Total MRR", "~$61,400/mo", "~$93,000\u2013$100,000/mo"],
    ],
    col_widths=[2.5, 1.8, 2.0]
)
add_callout(
    "$61,400/month self-funded by Year 2 is a strong, real business. "
    "$100,000/month requires outside capital or an unexpected organic acceleration. "
    "Both outcomes are worth pursuing."
)


# =============================================================================
# 8. TOTAL COST SUMMARY
# =============================================================================
add_h1("8. Total Cost to Launch \u2014 18-Month Summary")
add_table(
    ["Phase", "Timeframe", "Cost"],
    [
        ["Phase 0: Production ready + Unit compliance", "Month 1\u20132", "$5,000\u2013$10,000"],
        ["Phase 1: Website builder paid ads", "Month 3\u20136", "$20,000\u2013$30,000"],
        ["Phase 2: Mobile app + banking dev + compliance + ads", "Month 7\u201312", "$50,000\u2013$75,000"],
        ["Phase 3: Advisor + referral program dev + scaled ads", "Month 13\u201318", "$75,000\u2013$100,000"],
        ["Total 18-month runway needed", "", "$150,000\u2013$215,000"],
    ],
    col_widths=[3.2, 1.5, 1.6]
)
add_body(
    "Self-funding: aim for $150K in available runway before Month 3. "
    "Raising: a pre-seed round of $300,000\u2013$500,000 enables the $100K/month timeline."
)
add_h2("Advertising Spend Ramp")
add_table(
    ["Period", "Monthly Ad Spend", "Purpose"],
    [
        ["Month 3\u20136", "$5,000\u2013$8,000", "Find what converts. Test Meta, Google, YouTube. No scaling yet."],
        ["Month 7\u201312", "$8,000\u2013$12,000", "Scale winning creatives. Layer in mobile acquisition."],
        ["Month 13\u201318", "$12,000\u2013$18,000", "Scale behind proven CAC. Add advisor / banking leader campaigns."],
        ["Month 19\u201324 (with raise)", "$25,000\u2013$40,000", "Aggressive acquisition with referral flywheel offsetting CAC."],
    ],
    col_widths=[1.5, 1.8, 3.0]
)


# =============================================================================
# 9. THE FOUR DECISIONS
# =============================================================================
add_h1("9. The Four Decisions That Make or Break This")
add_h2("Decision 1 \u2014 Start Unit Production Application Now")
add_body(
    "Banking is your highest-revenue-potential stream and has the longest compliance lead time. "
    "Every week you delay is a week your banking launch slides. Submit the application in Month 1 "
    "regardless of where the rest of the product stands."
)
add_h2("Decision 2 \u2014 Build Onboarding Emails Before Running a Single Ad")
add_body(
    "If trials cannot find their way to a published website in 7 days, your conversion rate will be "
    "3\u20135% instead of 12\u201315%, and every ad dollar is worth half what it could be. "
    "The onboarding sequence (Day 1, 3, 7, 12, 14) must be live before Month 3."
)
add_h2("Decision 3 \u2014 First 20 Paying Orgs Must Come From Direct Outreach, Not Ads")
add_body(
    "Founder-led sales with churches in your personal and professional network produce testimonials, "
    "real product feedback, and word-of-mouth that no ad can replicate. "
    "Do this in Month 2\u20133 before you start paying for traffic."
)
add_h2("Decision 4 \u2014 Do Not Build the Advisor Network Until 100+ Paying Orgs")
add_body(
    "Advisors only pay $99\u2013$150/month if there is a real, active audience on the platform. "
    "Launching too early damages the feature's credibility. Target Month 13, not earlier."
)


# =============================================================================
# 10. RISKS
# =============================================================================
add_h1("10. Risks and Honest Mitigations")
add_table(
    ["Risk", "Likelihood", "Mitigation"],
    [
        ["Unit production approval delayed beyond Month 6", "Medium", "Start application Month 1. Hire fintech compliance consultant immediately."],
        ["Trial-to-paid conversion stays below 8%", "Medium", "Invest in onboarding email sequence before running ads. A/B test Day 7 and Day 12 nudges."],
        ["CAC exceeds $200 per trial on Meta", "Medium", "Diversify \u2014 podcast sponsorships, SEO, founder outreach. Don\u2019t bet everything on Meta."],
        ["Advisor network fails to attract sign-ups", "Low (if timed right)", "Don\u2019t launch until 100+ paying orgs. Founding cohort at $49/mo to de-risk."],
        ["Mobile app App Store rejection", "Low\u2013Medium", "Submit early, build in 4-week buffer for review cycles."],
        ["Competitor (Pushpay, Givelify) responds with pricing pressure", "Low", "Your 1% fee is already at the floor. Compete on website builder \u2014 they don\u2019t have it."],
    ],
    col_widths=[2.2, 1.2, 2.9]
)


# =============================================================================
# 11. SUMMARY
# =============================================================================
add_h1("11. Summary \u2014 The Path Is Clear")
add_body(
    "This is a real, fundable business with a complete technical foundation. "
    "Most startups spend 12\u201318 months building what you have already built. "
    "The remaining work is product polish, compliance, and disciplined distribution."
)
add_table(
    ["Milestone", "Target Date", "Revenue"],
    [
        ["Website builder live + first paid ads", "Month 3", "~$0 MRR"],
        ["50 paying orgs", "Month 6", "~$1,400/mo"],
        ["Mobile app + BankGO production launch", "Month 9\u201312", "~$8,000/mo"],
        ["Advisor network live", "Month 13", "~$10,000/mo"],
        ["All 5 revenue streams active", "Month 15", "~$18,000/mo"],
        ["Year 2 self-funded", "Month 24", "~$61,000/mo"],
        ["Year 2 with $300K raise", "Month 24", "~$100,000/mo"],
    ],
    col_widths=[2.8, 1.5, 2.0]
)
add_callout(
    "The website builder brings cost down.  "
    "The social platform brings referrals up.  "
    "Banking leaders bring BankGO adoption forward.  "
    "The advisor network makes the whole ecosystem more valuable to every participant.  "
    "That is the flywheel."
)

p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_end.paragraph_format.space_before = Pt(20)
r_end = p_end.add_run("Exchange Banking  \u00b7  Confidential Growth Plan  \u00b7  March 2026")
r_end.font.size = Pt(8.5)
r_end.font.color.rgb = SLATE
r_end.italic = True


# =============================================================================
# SAVE
# =============================================================================
out = "/Users/christopherfigurestest/Documents/give-app-main/docs/Exchange_Growth_Plan_2026.docx"
doc.save(out)
print(f"Saved: {out}")
